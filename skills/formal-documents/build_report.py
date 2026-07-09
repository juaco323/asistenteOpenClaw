#!/usr/bin/env python3
"""Genera informes .docx con la estructura fija §1-§7 de `formal-documents/SKILL.md`.

Uso:
    python3 build_report.py --input datos.json --output /ruta/informe.docx

El JSON de entrada declara "title" y "sections", con una clave por cada apartado
de SECTION_HEADINGS. Cada sección es un objeto con:
  - "body": texto del apartado. Párrafos separados por línea en blanco; una
    línea que empiece con "• " se convierte en viñeta.
  - "code" (opcional): lista de strings, cada una un bloque de código monoespaciado.
  - "references" (solo sección "7"): lista de referencias APA 7 (string por fuente).

Si falta una sección obligatoria o queda sin contenido, el script se detiene con
un error explícito en vez de generar un .docx incompleto: quien redacta el
informe debe escribir el texto o, si el apartado no aplica, una frase explícita
de no aplicabilidad (ver SKILL.md §3) antes de volver a generar el archivo.
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

SECTION_HEADINGS = {
    "1": "Resumen Ejecutivo",
    "2": "Introducción",
    "2.1": "Objetivos del Informe",
    "2.2": "Alcance y Limitaciones",
    "3": "Marco Conceptual",
    "4": "Desarrollo del Proyecto / Metodología",
    "4.1": "Arquitectura y Componentes",
    "4.2": "Implementación Técnica",
    "5": "Análisis de Resultados y Discusión",
    "6": "Conclusiones y Recomendaciones",
    "6.1": "Conclusiones",
    "6.2": "Recomendaciones y Trabajo Futuro",
    "7": "Referencias Bibliográficas",
}
SECTION_ORDER = list(SECTION_HEADINGS)
TOP_LEVEL = {"1", "2", "3", "4", "5", "6", "7"}

BODY_FONT = "Calibri"
CODE_FONT = "Consolas"


def _add_heading(doc: Document, number: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {SECTION_HEADINGS[number]}")
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(14 if number in TOP_LEVEL else 12)


def _add_body(doc: Document, body: str) -> None:
    for block in body.strip("\n").split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("• "):
            for line in block.split("\n"):
                line = line.lstrip("•").strip()
                if not line:
                    continue
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(line)
                run.font.name = BODY_FONT
                run.font.size = Pt(11)
        else:
            p = doc.add_paragraph()
            run = p.add_run(block)
            run.font.name = BODY_FONT
            run.font.size = Pt(11)
            p.paragraph_format.line_spacing = 1.15


def _add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = CODE_FONT
    run.font.size = Pt(10)


def _validate(data: dict) -> list[str]:
    errors = []
    if not str(data.get("title", "")).strip():
        errors.append("Falta 'title'.")
    sections = data.get("sections", {})
    missing = [n for n in SECTION_ORDER if n not in sections]
    if missing:
        errors.append("Faltan secciones: " + ", ".join(missing))
        return errors
    for number in SECTION_ORDER:
        section = sections[number]
        has_body = bool(str(section.get("body", "")).strip())
        if number == "7":
            has_refs = bool(section.get("references"))
            if not has_body and not has_refs:
                errors.append(
                    "Sección 7 (Referencias Bibliográficas) sin contenido: agrega "
                    "'references' (si se usó web/datos externos) o un 'body' que "
                    "indique explícitamente que no se usaron fuentes externas."
                )
        elif not has_body:
            errors.append(
                f"Sección {number} ({SECTION_HEADINGS[number]}) sin 'body'. "
                "Escribe el contenido o una frase explícita de no aplicabilidad."
            )
    return errors


def build(data: dict, output: Path) -> None:
    errors = _validate(data)
    if errors:
        raise ValueError("\n".join(errors))

    sections = data["sections"]
    doc = Document()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(data["title"].strip())
    title_run.bold = True
    title_run.font.name = BODY_FONT
    title_run.font.size = Pt(17)

    for number in SECTION_ORDER:
        section = sections[number]
        _add_heading(doc, number)
        if number == "7" and section.get("references"):
            for ref in section["references"]:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(str(ref))
                run.font.name = BODY_FONT
                run.font.size = Pt(11)
            if str(section.get("body", "")).strip():
                _add_body(doc, str(section["body"]))
        else:
            _add_body(doc, str(section.get("body", "")))
            for block in section.get("code") or []:
                _add_code(doc, str(block))

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"Informe guardado en: {output}")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path, help="JSON con title + sections")
    parser.add_argument("--output", required=True, type=Path, help="Ruta absoluta del .docx a generar")
    args = parser.parse_args()

    data = json.loads(args.input.read_text(encoding="utf-8"))
    try:
        build(data, args.output)
    except ValueError as exc:
        print(f"ERROR: estructura incompleta, no se generó el .docx.\n{exc}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
